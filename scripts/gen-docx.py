#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera docs/informe/ENTREGA_FINAL_SGROAS.docx (entrega final v1.0.0).

Uso: python scripts/gen-docx.py
Dependencia: python-docx (pip install python-docx)
"""
from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "informe" / "ENTREGA_FINAL_SGROAS.docx"

UTEQ = "UNIVERSIDAD TECNICA ESTATAL DE QUEVEDO"
TITULO = ("Sistema de Gestion de Recursos Operativos, Administrativos y de Seguridad "
          "(SGROAS) - Entrega Final v1.0.0")
GRUPO = ["CASTRO ESPINOZA KEVIN MOISES", "ESCUDERO PLAZA MARIA DEL ROSARIO",
         "TEJADA BAJANA LUIS ALEJANDRO"]
DOCENTE = "DR. GLEISTON CICERON GUERRERO ULLOA, PH.D."
CURSO = "5TO SOFTWARE A"
FECHA = "17 de agosto de 2026"
GITHUB = "https://github.com/Alxjandr07/SGROAS-ProyectoAppWeb.git"
DOI = "10.5281/zenodo.21698129"
GHCR = "ghcr.io/alxjandr07/sgroas"


def git(cmd: list[str]) -> str:
    try:
        return subprocess.run(["git", *cmd], cwd=ROOT, capture_output=True,
                              text=True, check=True).stdout.strip()
    except subprocess.CalledProcessError:
        return "N/A"


def tag() -> str:
    h = git(["rev-parse", "--short", "HEAD"])
    t = git(["describe", "--tags", "--exact-match"])
    return f"{t} ({h})" if t != "N/A" else f"main {h}"


def center(p, bold: bool = False, size: int = 11):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = bold
        r.font.size = Pt(size)
    return p


def clave_valor(doc, clave: str, valor: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{clave}: ")
    r.bold = True
    p.add_run(valor)


def tabla(doc, headers: list[str], filas: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for f in filas:
        row = t.add_row()
        for i, v in enumerate(f):
            row.cells[i].text = v
    return t


def main() -> None:
    doc = Document()
    doc.core_properties.title = "Entrega Final SGROAS v1.0.0"
    doc.core_properties.author = "Grupo CET - 5to Software A"

    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

    # ---------------- Portada ----------------
    p = doc.add_paragraph()
    center(p, bold=True, size=18)
    p.add_run(UTEQ)
    doc.add_paragraph()
    p = doc.add_paragraph()
    center(p, bold=True, size=13)
    p.add_run(TITULO)
    doc.add_paragraph()
    clave_valor(doc, "GRUPO", ", ".join(GRUPO))
    clave_valor(doc, "DOCENTE", DOCENTE)
    clave_valor(doc, "CURSO", CURSO)
    clave_valor(doc, "FECHA", FECHA)
    clave_valor(doc, "REPOSITORIO", GITHUB)
    clave_valor(doc, "ETIQUETA", tag())
    clave_valor(doc, "DOI SOFTWARE", DOI)
    clave_valor(doc, "IMAGEN CONTENEDOR", GHCR)
    doc.add_page_break()

    # ---------------- 1. Resumen ----------------
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "SGROAS (Sistema de Gestion de Recursos Operativos, Administrativos y de "
        "Seguridad) es la entrega final del Proyecto Fin de Curso (PFC) de la "
        "Universidad Tecnica Estatal de Quevedo, desarrollado por el grupo CET. El "
        "sistema gestiona conductores, vehiculos, rutas, asignaciones e incidentes de "
        "una cooperativa de transporte, con reportes estadisticos implementados como "
        "stored procedures (PostgreSQL, invocacion JPA 2.1 con REFCURSOR), cache "
        "Redis, autenticacion JWT, despliegue en Render y publicacion de imagen de "
        "contenedor en GitHub Container Registry (GHCR)."
    )
    doc.add_paragraph(
        "La entrega consolida: 165 pruebas automaticas con 0 fallos, cobertura "
        "JaCoCo 98,8 % (instrucciones) y 85,4 % (ramas), integracion continua con 3 "
        "corridas exitosas consecutivas, auditorias de SQL dinamico y trazabilidad "
        "aprobadas, mediciones de rendimiento k6 sin errores, Lighthouse 100/95/100/90 "
        "y usabilidad SUS 63,0 (Bueno)."
    )

    # ---------------- 2. Repositorio y CI/CD ----------------
    doc.add_heading("2. Repositorio y entrega continua", level=1)
    doc.add_paragraph(
        "El codigo fuente reside en GitHub y toda la linea de integracion continua se "
        "ejecuta sobre el flujo de trabajo .github/workflows/ci.yml, cubriendo build, "
        "pruebas, cobertura, auditorias y versionado."
    )
    tabla(doc,
          ["Elemento", "Valor"],
          [["Repositorio", GITHUB],
           ["Rama principal", "main"],
           ["Corridas CI exitosas (entregables A6/A7)", "3 de 3 consecutivas"],
           ["Etiqueta de entrega", tag()],
           ["Imagen GHCR", GHCR + " (v1.0.0 y latest)"],
           ["DOI software", DOI]])
    doc.add_paragraph(
        "El job docker-publish se activa sobre etiquetas v* y publica la imagen "
        "multi-plataforma (buildx) con login autenticado por GITHUB_TOKEN, dejando el "
        "digest en la salida del job."
    )

    # ---------------- 3. Arquitectura ----------------
    doc.add_heading("3. Arquitectura y decisiones de diseno", level=1)
    doc.add_paragraph(
        "Backend Spring Boot con persistencia hibrida CRUD-ORM (JPA) y acceso a "
        "datos por stored procedures para reportes (ADR-006); API REST con DTOs "
        "(ADR-005b); despliegue PaaS en Render con Postgres y Redis administrados "
        "(ADR-007). Seguridad JWT (RFC 7519) con cookies HttpOnly y rate limiting de "
        "login."
    )
    doc.add_paragraph("Registro de decisiones (docs/adr/):")
    for adr in sorted((ROOT / "docs" / "adr").glob("adr-*.md")):
        doc.add_paragraph(f"- {adr.name}", style="List Bullet")
    doc.add_paragraph(
        "Los 7 reportes estadisticos se implementan como PROCEDURE con REFCURSOR y "
        "se invocan via @NamedStoredProcedureQuery + @Procedure (JPA 2.1), cubiertos "
        "por StoredProcedureIntegrationTest con 7 pruebas sobre PostgreSQL real."
    )

    # ---------------- 4. Verificacion ----------------
    doc.add_heading("4. Verificacion y calidad", level=1)
    tabla(doc,
          ["Criterio", "Resultado"],
          [["Pruebas JUnit 5 (./mvnw clean verify)", "165 ejecutadas, 0 fallos, 0 errores"],
           ["Cobertura JaCoCo - instrucciones", "98,8 % (umbral >= 60 %)"],
           ["Cobertura JaCoCo - ramas", "85,4 % (umbral >= 60 %)"],
           ["Cobertura JaCoCo - lineas", "99,7 %"],
           ["SpotBugs", "0 bugs (BugInstance size is 0)"],
           ["Auditoria SQL dinamico (audit-sql-dynamic.sh)", "0 violaciones"],
           ["Auditoria trazabilidad (validate-traceability.sh)", "matriz valida (19 requisitos)"],
           ["Despliegue local", "docker compose (PostgreSQL + Redis + backend)"]])

    # ---------------- 5. Mediciones ----------------
    doc.add_heading("5. Mediciones empiricas", level=1)
    doc.add_heading("5.1 Rendimiento (k6, Bloque C.1)", level=2)
    doc.add_paragraph(
        "3 corridas de 50 VUs / 30 s sobre GET /api/conductores con cache caliente: "
        "4.468 peticiones totales, verificaciones 8.930 exitosas y 0 fallidas, tasa de "
        "error 0,00 %, p95 medio 81,43 ms (maximo 173,15 ms, umbral < 200 ms)."
    )
    tabla(doc,
          ["Corrida", "Reqs", "avg (ms)", "p95 (ms)", "p99 (ms)", "Errores"],
          [["1 (k01-run1.json)", "1466", "37,26", "173,15", "251,54", "0"],
           ["2 (k02-run2.json)", "1501", "16,82", "34,82", "155,99", "0"],
           ["3 (k03-run3.json)", "1501", "14,70", "36,32", "156,97", "0"]])
    doc.add_heading("5.2 Accesibilidad y calidad web (Lighthouse, Bloque C.5)", level=2)
    tabla(doc,
          ["Corrida", "Performance", "Accessibility", "Best Practices", "SEO"],
          [["1 (lhci-20260730-2115)", "100", "95", "100", "90"],
           ["2 (lhci-20260730-2117)", "100", "95", "100", "90"]])
    doc.add_heading("5.3 Usabilidad (SUS, Bloque C.3)", level=2)
    doc.add_paragraph(
        "10 participantes (P01 a P10): media 63,0 / 100 (DT 13,88, IC 95 % "
        "[53,07; 72,93]), calificacion adjetiva 'Bueno' (Bangor et al., 2009)."
    )

    # ---------------- 6. Despliegue ----------------
    doc.add_heading("6. Despliegue y operacion", level=1)
    doc.add_paragraph(
        "Despliegue PaaS en Render mediante blueprint render.yaml (Postgres "
        "administrado, Redis y servicio web con health check /actuator/health). "
        "Documentacion operativa en docs/despliegue/: DEPLOYMENT.md, RUNBOOK.md y "
        "BACKUP.md, con script de respaldo periodico scripts/backup-prod.sh."
    )

    # ---------------- 7. Entregables ----------------
    doc.add_heading("7. Entregables de la entrega final", level=1)
    tabla(doc,
          ["Artefacto", "Ubicacion"],
          [["Informe academico LaTeX", "docs/informe/main.tex"],
           ["Documento de despliegue", "docs/despliegue/DEPLOYMENT.md"],
           ["Runbook de operacion", "docs/despliegue/RUNBOOK.md"],
           ["Respaldos", "docs/despliegue/BACKUP.md"],
           ["ADRs (001-007)", "docs/adr/"],
           ["Mediciones (jacoco/perf/lighthouse/sus/sec)", "docs/mediciones/"],
           ["Observaciones y evidencias", "docs/observaciones/, docs/etica/"],
           ["Esquema y datos de base", "docs/basedatos/, sgroas_bd_esquema.sql"],
           ["Makefile pipeline (make all)", "Makefile"],
           ["Este documento", "docs/informe/ENTREGA_FINAL_SGROAS.docx"]])

    doc.add_paragraph(
        f"Documento generado automaticamente por scripts/gen-docx.py el "
        f"{date.today().isoformat()} sobre la etiqueta {tag()}."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()